import re
from docx import Document
from PIL import Image
import fitz
from nltk.tokenize import sent_tokenize

def process_txt(file_path: str):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

def preprocess_text(text: str):
    return clean_text(text)

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def split_into_chunks(text: str, chunk_size: int, overlap: int) -> List[str]:
    """
    Разделение текста на контекстные окна фиксированного размера.

    :param text: Исходный текст.
    :param chunk_size: Размер окна.
    :param overlap: Количество перекрывающихся токенов между окнами.
    :return: Список текстовых кусочков.
    """
    sentences = re.split(r"(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s", text)
    chunks = []
    current_chunk = []

    current_length = 0
    for sentence in sentences:
        sentence_length = len(sentence.split())
        if current_length + sentence_length > chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = current_chunk[-overlap:]
            current_length = len(" ".join(current_chunk).split())

        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks







# def process_doc(file_path: str):
#     doc = Document(file_path)
#     text = "\n".join([p.text for p in doc.paragraphs])
#     images = extract_images_from_doc(file_path)
#     return {
#         "text": text,
#         "images": describe_images(images),
#         "structure": DoclingProcessor.analyze(text)
#     }

# def process_pdf(file_path: str):
#     doc = fitz.open(file_path)
#     text = ""
#     images = []
#     for page in doc:
#         text += page.get_text()
#         images.extend(extract_images_from_pdf(file_path))
#     return {
#         "text": text,
#         "images": describe_images(images),
#         "structure": ColpaliModel.analyze(text)
#     }


# def process_pdf(file_path: str):
#     doc = fitz.open(file_path)
#     text = ""
#     images = []
#     for page in doc:
#         text += page.get_text()
#         images.extend(page.get_images(full=True))
#     return {
#         "text": text,
#         "images": describe_images(images),
#         "structure": ColpaliModel.analyze(text)
#     }

# def extract_images_from_doc(file_path):
#     """
#     Извлекает изображения из файла .docx или .doc.
#     :param file_path: Путь к файлу.
#     :return: Список изображений (в виде PIL Image).
#     """
#     doc = Document(file_path)
#     images = []
#     for rel in doc.part.rels.values():
#         if "image" in rel.target_ref:
#             image_data = rel.target_part.blob
#             image = Image.open(io.BytesIO(image_data))
#             images.append(image)
#     return images

# def extract_images_from_pdf(file_path):
#     """
#     Извлекает изображения из PDF.
#     :param file_path: Путь к файлу.
#     :return: Список изображений (в виде PIL Image).
#     """
#     pdf = fitz.open(file_path)
#     images = []
#     for page_num in range(len(pdf)):
#         page = pdf[page_num]
#         for img_index, img in enumerate(page.get_images(full=True)):
#             xref = img[0]
#             base_image = pdf.extract_image(xref)
#             image_bytes = base_image["image"]
#             image = Image.open(io.BytesIO(image_bytes))
#             images.append(image)
#     return images
